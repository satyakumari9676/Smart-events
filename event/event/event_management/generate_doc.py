import os
import docx
from docx.shared import Pt
import time

def create_doc():
    doc = docx.Document()
    doc.add_heading('Event Management System - Source Code', 0)

    files_to_include = [
        'event_management/settings.py',
        'event_management/urls.py',
        'accounts/models.py',
        'accounts/views.py',
        'accounts/urls.py',
        'accounts/admin.py',
        'pages/views.py',
        'pages/urls.py',
        'pages/models.py',
        'budget/views.py',
        'budget/urls.py',
        'templates/base.html',
        'templates/index.html',
        'templates/staff_dashboard.html',
        'templates/client_dashboard.html',
        'templates/worker_dashboard.html',
        'templates/login.html'
    ]

    for file_path in files_to_include:
        if os.path.exists(file_path):
            try:
                # Add File Header
                doc.add_heading(f'File: {file_path}', level=1)
                
                # Add File Content
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()

                # Add content as code paragraph
                p = doc.add_paragraph()
                run = p.add_run(content)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                
                doc.add_page_break()
                print(f"Added {file_path}")
            except Exception as e:
                print(f"Failed to read {file_path}: {e}")
        else:
            print(f"File {file_path} not found.")

    output_path = r'D:\event\Project_Source_Code.docx'
    doc.save(output_path)
    print(f"\nDocument successfully created at: {output_path}")

if __name__ == '__main__':
    create_doc()
